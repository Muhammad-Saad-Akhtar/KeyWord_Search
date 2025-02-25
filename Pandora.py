import os
from docx import Document
from docx.shared import RGBColor
from tqdm import tqdm  # For progress bar

def search_in_docx(folder_path, search_text, exclude_words=[], progress_callback=None):
    """
    Searches for text in .docx files within a given folder.

    Args:
        folder_path (str): Path to the folder containing .docx files.
        search_text (str): The text to search for.
        exclude_words (list): Words to exclude from results.
        progress_callback (function, optional): Function to update progress (files_scanned, total_files).

    Returns:
        list: A list of tuples (filename, paragraph_number, found_text).
    """
    if not os.path.isdir(folder_path):
        raise FileNotFoundError(f"Invalid folder path: {folder_path}")

    results = []
    docx_files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
    
    if not docx_files:
        print("No .docx files found in the folder.")
        return results

    total_files = len(docx_files)

    for idx, filename in enumerate(tqdm(docx_files, desc="Scanning files", unit="file"), start=1):
        file_path = os.path.join(folder_path, filename)

        try:
            doc = Document(file_path)
            for i, para in enumerate(doc.paragraphs):
                text = para.text.lower()
                if search_text.lower() in text and not any(word.lower() in text for word in exclude_words):
                    results.append((filename, i + 1, para.text))
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            results.append((filename, None, f"Error reading file: {e}"))

        if progress_callback:
            progress_callback(idx, total_files)

    return results


def highlight_and_save_results(results, search_text):
    """
    Highlights the found search text in the document and saves it as a new file.

    Args:
        results (list): The list of search results containing filename, paragraph number, and found text.
        search_text (str): The text to search for and highlight.
    """
    source_folder = r"C:\Users\HP\Desktop\Others\Keyword Searcher\The guardians of Pandora"
    output_folder = r"C:\Users\HP\Desktop\New-Pandora"

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename, para_number, found_text in results:
        file_path = os.path.join(source_folder, filename)
        output_path = os.path.join(output_folder, f"highlighted_{filename}")

        try:
            doc = Document(file_path)

            for para in doc.paragraphs:
                if search_text.lower() in para.text.lower():
                    for run in para.runs:
                        if search_text.lower() in run.text.lower():
                            run.font.highlight_color = 6  # Yellow highlight (WORD_HIGHLIGHT_YELLOW)

            doc.save(output_path)
            print(f"Saved highlighted version of {filename} in {output_folder}.")
        
        except Exception as e:
            print(f"Error processing {filename}: {e}")


if __name__ == "__main__":
    folder = r"C:\Users\HP\Desktop\Others\Keyword Searcher\KeyWord_Searcher\The guardians of Pandora"

    while True:
        text = input("Enter text to search (press 'e' to exit): ")
        if text.lower() == 'e':
            print("Exiting program.")
            break

        results = search_in_docx(folder, text, exclude_words=["example", "sample"])

        if results:
            print("\nSearch Results:")
            for filename, line_number, found_text in results:
                if line_number:
                    print(f"Found in {filename}, Paragraph {line_number}: {found_text}")
                else:
                    print(f"{filename}: {found_text}")

            highlight_and_save_results(results, text)
        else:
            print("No matches found.")
