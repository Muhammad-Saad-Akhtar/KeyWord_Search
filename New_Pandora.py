import os
import docx
from docx import Document

def search_in_docx(folder_path, search_text, exclude_words=[]):
    if not os.path.isdir(folder_path):
        print("Invalid folder path.")
        return
    
    results = []
    total_files = len([f for f in os.listdir(folder_path) if f.endswith(".docx")])
    
    output_folder = r"C:\\Users\\HP\\Desktop\\New-Pandora"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    for idx, filename in enumerate(os.listdir(folder_path), start=1):
        if not filename.endswith(".docx"): 
            continue
        
        file_path = os.path.join(folder_path, filename)
        output_file_path = os.path.join(output_folder, f"Highlighted_{filename}")
        
        try:
            doc = Document(file_path)
            new_doc = Document()
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.lower()
                if search_text.lower() in text and not any(word in text for word in exclude_words):
                    highlighted_text = para.text.replace(search_text, f"**{search_text}**")
                    new_doc.add_paragraph(highlighted_text)
                    results.append((filename, i + 1, para.text))
                else:
                    new_doc.add_paragraph(para.text)
            
            new_doc.save(output_file_path)
            print(f"Progress: {idx}/{total_files} files scanned")
        except Exception as e:
            print(f"Could not read {filename}: {e}")
    
    if results:
        print("\nSearch Results:")
        for filename, line_number, found_text in results:
            highlighted_text = found_text.replace(search_text, f"**{search_text}**")
            print(f"Found in {filename}, Paragraph {line_number}: {highlighted_text}")
    else:
        print("No matches found.")
    
if __name__ == "__main__":
    folder = r"C:\\Users\\HP\\Desktop\\The guardians of Pandora"
    
    while True:
        text = input("Enter text to search (press 'e' to exit): ")
        if text.lower() == 'e':
            print("Exiting program.")
            break
        search_in_docx(folder, text, exclude_words=["example", "sample"])
